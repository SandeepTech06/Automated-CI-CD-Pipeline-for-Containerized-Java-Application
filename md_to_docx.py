from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import re

md_path = r"s:\Automated CICD Pipeline\SYNOPSIS.md"
docx_path = r"s:\Automated CICD Pipeline\SYNOPSIS.docx"

def add_paragraph_with_style(doc, text, style_name=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = Document()
styles = doc.styles
# ensure Normal font size
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

list_mode = False
for raw in lines:
    line = raw.rstrip('\n')
    if not line.strip():
        # blank line -> paragraph break
        doc.add_paragraph()
        list_mode = False
        continue
    # headings
    if line.startswith('#'):
        hashes = len(re.match(r"^#+", line).group(0))
        text = line[hashes:].strip()
        level = min(hashes, 3)
        doc.add_heading(text, level=level)
        list_mode = False
        continue
    # unordered list
    m = re.match(r"^[-*+]\s+(.*)", line)
    if m:
        text = m.group(1)
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(text)
        list_mode = True
        continue
    # numbered list
    m = re.match(r"^\d+[.)]\s+(.*)", line)
    if m:
        text = m.group(1)
        p = doc.add_paragraph(style='List Number')
        p.add_run(text)
        list_mode = True
        continue
    # code block fence (skip triple backticks) - represent as preformatted paragraph
    if line.strip().startswith('```'):
        # toggle code block - naive: collect until next ```
        code_lines = []
        # read following lines until ```
        continue
    # normal paragraph
    doc.add_paragraph(line)

# Save document
try:
    doc.save(docx_path)
    print(f"Saved: {docx_path}")
except Exception as e:
    print("Error saving docx:", e)
